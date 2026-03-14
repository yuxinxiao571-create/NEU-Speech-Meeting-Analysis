# ===================== 先导入config.py =====================
from config import *
import os

# ===================== 替换原代码中的硬编码配置 =====================
# 1. 设置 HuggingFace 镜像站（解决网络问题）
os.environ["HF_ENDPOINT"] = HF_ENDPOINT
os.environ["HF_HUB_ENABLE_HF_TRANSFER"] = "1"  # 启用高速下载

# 2. 设置模型缓存根目录
os.environ["HF_HOME"] = HF_CACHE_DIR
os.environ["TRANSFORMERS_CACHE"] = os.path.join(HF_CACHE_DIR, "transformers")
os.environ["HUGGINGFACE_HUB_CACHE"] = os.path.join(HF_CACHE_DIR, "hub")
os.environ["SENTENCE_TRANSFORMERS_HOME"] = os.path.join(HF_CACHE_DIR, "sentence-transformers")

# 3. 强制离线模式（关键！）
os.environ["HF_HUB_OFFLINE"] = "1"

# 创建缓存目录（如果不存在）
os.makedirs(HF_CACHE_DIR, exist_ok=True)

# 解决 tokenizers 并行警告
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# 4. 替换所有硬编码配置
device = DEVICE
SEMANTIC_MODEL_PATH = SEMANTIC_MODEL_PATH
DEEPSEEK_API_KEY = DEEPSEEK_API_KEY
MODEL_NAME = MODEL_NAME
audio_dir = AUDIO_DIR
output_dir = OUTPUT_DIR
integrated_dir = INTEGRATED_DIR
GENERATION_ACCURACY_THRESHOLD = GENERATION_ACCURACY_THRESHOLD
MIN_CLUSTER_SIZE = MIN_CLUSTER_SIZE
MIN_SOURCE_THRESHOLD = MIN_SOURCE_THRESHOLD
ALLOW_GENERAL_SOURCE = ALLOW_GENERAL_SOURCE
MAX_DOWNLOAD_RETRIES = MAX_DOWNLOAD_RETRIES

# ===================== 其余代码保持不变 =====================
# 忽略 pyannote 的 std() 警告
import warnings
warnings.filterwarnings("ignore", message="std(): degrees of freedom is <= 0.")

import torch
from pathlib import Path
from pyannote.core import Segment, Annotation
import whisper
import time
import re
import requests
import json
from pyannote.audio import Pipeline
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from openpyxl.styles import Font, Alignment, PatternFill
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from collections import defaultdict
from scipy.cluster.hierarchy import linkage, fcluster
from sentence_transformers import SentenceTransformer
from huggingface_hub import snapshot_download
from sklearn.metrics import precision_score, recall_score, f1_score
import random


# 发言人映射（保持中文）
SPEAKER_MAPPING = {
    "SPEAKER_00": "张教授（电力系统专家）",
    "SPEAKER_01": "李工程师（充电桩厂商）",
    "SPEAKER_02": "王主任（学校后勤）",
}


# -------------------------- 工具函数 --------------------------
def delete_old_files(output_dir, prefix=None):
    if not os.path.exists(output_dir):
        return
    deleted_count = 0
    for filename in os.listdir(output_dir):
        file_path = os.path.join(output_dir, filename)
        if os.path.isfile(file_path) and (prefix is None or filename.startswith(prefix)):
            if filename.lower().endswith((".docx", ".xlsx")):
                os.remove(file_path)
                deleted_count += 1


def sanitize_text(text):
    if not text:
        return ""
    text = str(text)
    control_chars = re.compile(r'[\x00-\x1F\x7F-\x9F]')
    text = control_chars.sub('', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


# -------------------------- 1. 语音转文本与说话人分离 --------------------------
def get_text_with_timestamp(transcribe_res):
    timestamp_texts = []
    for item in transcribe_res["segments"]:
        start = item["start"]
        end = item["end"]
        text = item["text"]
        timestamp_texts.append((Segment(start, end), text))
    return timestamp_texts


def add_speaker_info_to_text(timestamp_texts, ann):
    spk_text = []
    for seg, text in timestamp_texts:
        spk = ann.crop(seg).argmax()
        spk_text.append((seg, spk, text))
    return spk_text


def merge_cache(text_cache):
    sentence = ''.join([item[-1] for item in text_cache])
    spk = text_cache[0][1]
    start = round(text_cache[0][0].start, 1)
    end = round(text_cache[-1][0].end, 1)
    mapped_spk = SPEAKER_MAPPING.get(spk, spk)
    return Segment(start, end), spk, mapped_spk, sentence, start, end


PUNC_SENT_END = ['.', '?', '!', "。", "？", "！"]
def merge_sentence(spk_text):
    merged_spk_text = []
    pre_spk = None
    text_cache = []
    for seg, spk, text in spk_text:
        if (spk != pre_spk and len(text_cache) > 0) or not text.strip():
            if text_cache:
                merged_seg, raw_spk, mapped_spk, merged_sent, start, end = merge_cache(text_cache)
                merged_spk_text.append((merged_seg, raw_spk, mapped_spk, merged_sent, start, end))
            text_cache = [(seg, spk, text)] if text.strip() else []
            pre_spk = spk
        else:
            text_cache.append((seg, spk, text))
            pre_spk = spk
            if any(text.endswith(punc) for punc in PUNC_SENT_END):
                merged_seg, raw_spk, mapped_spk, merged_sent, start, end = merge_cache(text_cache)
                merged_spk_text.append((merged_seg, raw_spk, mapped_spk, merged_sent, start, end))
                text_cache = []
                pre_spk = None
    if len(text_cache) > 0:
        merged_seg, raw_spk, mapped_spk, merged_sent, start, end = merge_cache(text_cache)
        merged_spk_text.append((merged_seg, raw_spk, mapped_spk, merged_sent, start, end))
    return merged_spk_text


def diarize_text(transcribe_res, diarization_result):
    timestamp_texts = get_text_with_timestamp(transcribe_res)
    spk_text = add_speaker_info_to_text(timestamp_texts, diarization_result)
    return merge_sentence(spk_text)


# -------------------------- 2. 大模型处理 --------------------------
def format_dialogue_for_llm(dialogue):
    llm_input = "以下是专家讨论的逐句记录（含发言人和精确时间戳），请严格基于这些内容生成结构化分析：\n"
    for i, item in enumerate(dialogue, 1):
        spk = item['mapped_speaker']
        start = item['start']
        end = item['end']
        text = sanitize_text(item['text'])
        llm_input += f"[{i}] [{start}s-{end}s] {spk}：{text}\n"
    return llm_input


def generate_structured_with_llm(llm_input):
    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
    }
    prompt = f"""
    任务：基于以下讨论记录生成结构化分析，核心要求是「100%信息保真」。
    强制规则：
    1. 内容绝对保真：结论必须是原始对话的直接提取，保留关键动词/名词（如“成本控制在10万内”而非“成本较低”），禁止添加主观解读。
    2. 来源强制绑定：每条结论必须标注来源，格式为“[序号] [时间戳] 发言人”（如[1] [10.2s-15.5s] 张教授），无序号的结论直接删除。
    3. 冲突优先标注：若存在对立观点（如A支持X、B反对X），需先标注“[冲突组]”，再分点列出（如2.1 支持X：...；2.2 反对X：...）。
    4. 空值处理：若某模块（如“讨论结果”）无对应内容，标注“无相关信息”，禁止编造。

    讨论记录：
    {llm_input}

    输出格式（严格遵守）：
    1. 议题：
    1.1 [核心议题]（来源：[序号] [时间戳] 发言人；[序号] [时间戳] 发言人）
    2. 专家意见：
    2.1 [冲突组] 支持观点：[原始表述]（来源：[序号] [时间戳] 发言人）
    2.2 [冲突组] 反对观点：[原始表述]（来源：[序号] [时间戳] 发言人）
    2.3 无冲突观点：[原始表述]（来源：[序号] [时间戳] 发言人）
    3. 讨论结果：
    3.1 [结论]（来源：[序号] [时间戳] 发言人）/ 无相关信息
    """

    data = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": "你是严谨的会议记录分析师，必须完整保留所有冲突观点并标注准确来源，严格遵守格式要求。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.1,
        "timeout": 60
    }

    for retry in range(3):
        try:
            response = requests.post(url, headers=headers, data=json.dumps(data), timeout=60)
            response.raise_for_status()
            response_data = response.json()
            if response_data.get("choices") and response_data["choices"][0]["message"]["content"].strip():
                return response_data["choices"][0]["message"]["content"].strip()
            print(f"大模型返回空结果，重试第{retry+1}次...")
        except Exception as e:
            print(f"大模型调用失败（{retry+1}/3）：{str(e)}")
            time.sleep(2)
    return "生成失败：大模型调用多次失败"


def parse_llm_output(structured_text, file_id, raw_dialogue=None):
    """优化解析逻辑：增加输出校验，过滤无效内容，提升准确性"""
    sections = {"议题": [], "专家意见": [], "讨论结果": []}
    current_section = None
    source_pattern = re.compile(r'（来源：(.*?)）$')
    index_pattern = re.compile(r'\[(\d+)\]')

    if raw_dialogue is None:
        raw_dialogue = []
    raw_texts = [sanitize_text(item["text"]).lower() for item in raw_dialogue]

    print(f"解析结构化文本：\n{structured_text}\n")

    for line in structured_text.split('\n'):
        line = line.strip()
        if not line or "无相关信息" in line:
            continue

        if line.startswith(("1. 议题", "议题：", "一、议题")):
            current_section = "议题"
            continue
        elif line.startswith(("2. 专家意见", "专家意见：", "二、专家意见")):
            current_section = "专家意见"
            continue
        elif line.startswith(("3. 讨论结果", "讨论结果：", "三、讨论结果")):
            current_section = "讨论结果"
            continue

        if current_section:
            source_match = source_pattern.search(line)
            if not source_match:
                print(f"过滤无来源内容：{line}")
                continue

            source_str = source_match.group(1).strip()
            source_indices = index_pattern.findall(source_str)
            valid_indices = []
            for idx in source_indices:
                try:
                    idx_int = int(idx) - 1
                    if 0 <= idx_int < len(raw_dialogue):
                        valid_indices.append(idx)
                except:
                    pass
            if not valid_indices:
                print(f"过滤无效来源内容：{line}")
                continue

            content = line[:source_match.start()].strip().lower()
            content_keywords = set(re.findall(r'\b\w{2,}\b', content)) - {"的", "是", "在", "应"}
            if not content_keywords:
                continue

            match_flag = False
            for idx in valid_indices:
                idx_int = int(idx) - 1
                raw_text = raw_texts[idx_int]
                if any(keyword in raw_text for keyword in content_keywords):
                    match_flag = True
                    break
            if not match_flag:
                print(f"过滤偏离原始内容：{line}")
                continue

            sections[current_section].append({
                "content": content,
                "sources": [f"[{idx}] {raw_dialogue[int(idx)-1]['start']}s-{raw_dialogue[int(idx)-1]['end']}s {raw_dialogue[int(idx)-1]['mapped_speaker']}" for idx in valid_indices],
                "file_id": file_id,
                "valid_source": True,
                "raw_content": line
            })

    if raw_dialogue and len(sections["专家意见"]) < 2:
        print(f"警告：专家意见不足，从原始对话补充观点（防止冲突丢失）")
        added_texts = set()
        for idx, item in enumerate(raw_dialogue[:8]):
            text = sanitize_text(item["text"]).split("：")[-1].strip()
            if text not in added_texts and len(text) > 5:
                sections["专家意见"].append({
                    "content": text,
                    "sources": [f"[{idx+1}] {item['start']}s-{item['end']}s {item['mapped_speaker']}"],
                    "file_id": file_id,
                    "raw_content": text,
                    "valid_source": True
                })
                added_texts.add(text)

    return sections


# -------------------------- 3. 语义级冲突检测模块 --------------------------
def adaptive_semantic_clustering(texts, semantic_model, min_cluster_size=2):
    """优化语义分组：关键词预过滤+组内相似度校验，减少无效分组"""
    if len(texts) <= 1:
        return [[0]] if texts else []

    def extract_core_keywords(text):
        topic_keywords = ["安装位置", "功率", "成本", "维护", "扩容", "负荷", "预算", "周期"]
        return [kw for kw in topic_keywords if kw in text]

    keyword_groups = defaultdict(list)
    for idx, text in enumerate(texts):
        keywords = extract_core_keywords(text)
        if keywords:
            key = "-".join(sorted(keywords))
            keyword_groups[key].append(idx)
        else:
            keyword_groups["other"].append(idx)

    final_groups = []
    for group_indices in keyword_groups.values():
        if len(group_indices) < min_cluster_size:
            final_groups.append(group_indices)
            continue

        group_texts = [texts[idx] for idx in group_indices]
        embeddings = semantic_model.encode(group_texts, normalize_embeddings=True)
        similarity_matrix = cosine_similarity(embeddings)
        distance_matrix = 1 - similarity_matrix

        Z = linkage(distance_matrix[np.triu_indices_from(distance_matrix, k=1)], method='ward')
        max_clusters = min(len(group_texts) // 2, 4)
        if max_clusters < 2:
            max_clusters = 2

        sse = []
        for n_clusters in range(1, max_clusters + 1):
            clusters = fcluster(Z, t=n_clusters, criterion='maxclust')
            cluster_labels = list(set(clusters))
            sse_total = 0.0
            for label in cluster_labels:
                c_indices = [i for i, c in enumerate(clusters) if c == label]
                if len(c_indices) < 1:
                    continue
                c_distances = distance_matrix[np.ix_(c_indices, c_indices)]
                sse_total += np.mean(c_distances)
            sse.append(sse_total)

        if len(sse) < 2:
            optimal_clusters = 1
        else:
            slopes = [sse[i] - sse[i+1] for i in range(len(sse)-1)]
            optimal_clusters = np.argmax(slopes) + 2

        clusters = fcluster(Z, t=optimal_clusters, criterion='maxclust')
        sub_group_dict = defaultdict(list)
        for sub_idx, cluster_id in enumerate(clusters):
            original_idx = group_indices[sub_idx]
            sub_group_dict[cluster_id].append(original_idx)

        for sub_group in sub_group_dict.values():
            if len(sub_group) < min_cluster_size:
                final_groups.append(sub_group)
                continue
            sub_embeddings = semantic_model.encode([texts[idx] for idx in sub_group], normalize_embeddings=True)
            avg_sim = np.mean(cosine_similarity(sub_embeddings))
            if avg_sim < 0.3:
                final_groups.extend([[idx] for idx in sub_group])
            else:
                final_groups.append(sub_group)

    remaining = [idx for group in final_groups if len(group) < min_cluster_size for idx in group]
    valid_groups = [group for group in final_groups if len(group) >= min_cluster_size]

    for idx in remaining:
        if not valid_groups:
            valid_groups.append([idx])
            continue
        idx_emb = semantic_model.encode([texts[idx]], normalize_embeddings=True)
        max_sim = -1
        best_group_idx = 0
        for g_idx, group in enumerate(valid_groups):
            group_emb = semantic_model.encode([texts[g_idx] for g_idx in group], normalize_embeddings=True)
            sim = cosine_similarity(idx_emb, group_emb).mean()
            if sim > max_sim:
                max_sim = sim
                best_group_idx = g_idx
        if max_sim >= 0.2:
            valid_groups[best_group_idx].append(idx)
        else:
            valid_groups.append([idx])

    return valid_groups


def enhanced_analyze_conflict_with_llm(texts, topic=None):
    """增强型冲突判定：明确冲突类型+反例对比，提升精准度"""
    url = "https://api.deepseek.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
    }

    core_topic = topic if topic else "未知议题"
    conflict_types = """
    常见冲突类型（仅以下情况判定为冲突）：
    1. 对立主张：如“支持A方案”vs“反对A方案”，“选X位置”vs“选Y位置”；
    2. 互斥结论：如“成本10万内”vs“成本需20万”，“每月维护”vs“每季度维护”；
    3. 否定关系：如“需扩容”vs“无需扩容”，“有风险”vs“无风险”。

    非冲突情况（即使有差异也不判定为冲突）：
    1. 细节补充：如“A方案成本10万”vs“A方案成本10万，含人工”；
    2. 角度差异：如“A方案优点是快”vs“A方案缺点是贵”；
    3. 数值接近：如“成本10万”vs“成本12万”（无主张对立）。
    """

    text_list = "\n".join([f"{i+1}. {t}" for i, t in enumerate(texts)])
    prompt = f"""
    任务：基于议题「{core_topic}」，判断以下观点是否存在冲突，严格参考冲突类型标准。

    {conflict_types}

    观点列表：
    {text_list}

    输出要求：
    1. 先明确“是否冲突”，再计算“冲突概率”（0.0-1.0，越确定越接近1）；
    2. “冲突点”需说明具体对立关系（如“安装位置：停车场vs教学楼”）；
    3. “判定依据”需对应上述冲突类型（如“属于对立主张，符合冲突类型1”）。

    仅返回JSON（无其他文字）：
    {{
        "has_conflict": true/false,
        "probability": 0.0-1.0,
        "details": "冲突点（无冲突则为null）",
        "reasoning": "判定依据（需关联冲突类型）"
    }}
    """

    data = {
        "model": MODEL_NAME,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.05,
        "timeout": 35
    }

    for retry in range(2):
        try:
            response = requests.post(url, headers=headers, data=json.dumps(data), timeout=35)
            response.raise_for_status()
            res_data = response.json()
            if res_data.get("choices"):
                llm_out = res_data["choices"][0]["message"]["content"].strip()
                json_match = re.search(r'\{.*\}', llm_out, re.DOTALL)
                if json_match:
                    result = json.loads(json_match.group())
                    result["probability"] = min(max(float(result.get("probability", 0.0)), 0.0), 1.0)
                    return result
        except Exception as e:
            print(f"冲突判定重试（{retry+1}/2）：{e}")
            time.sleep(1)

    conflict_keywords = ["反对", "不建议", "无需", "不能", "相反", "而是"]
    has_conflict = any(kw in text.lower() for text in texts for kw in conflict_keywords)
    return {
        "has_conflict": has_conflict,
        "probability": 0.6 if has_conflict else 0.1,
        "details": "关键词检测到潜在冲突" if has_conflict else None,
        "reasoning": "大模型调用失败，基于关键词兜底判定"
    }


def optimized_detect_semantic_conflicts(all_sections, semantic_model):
    """优化后的冲突检测主函数：整合自适应分组+增强判定+动态阈值"""
    integrated = defaultdict(list)
    conflict_logs = []

    for section_type, items in all_sections.items():
        if not items:
            continue
        print(f"处理{section_type}：共{len(items)}条观点")

        texts = [item["content"] for item in items if item["content"].strip()]
        valid_indices = [i for i, item in enumerate(items) if item["content"].strip()]
        if len(texts) < 2:
            integrated[section_type].append({
                "core_topic": texts[0] if texts else "无内容",
                "related_items": items,
                "has_conflict": False,
                "conflict_details": None,
                "conflict_probability": 0.0,
                "reasoning": "观点数量不足"
            })
            continue

        groups = adaptive_semantic_clustering(texts, semantic_model, min_cluster_size=MIN_CLUSTER_SIZE)
        print(f"{section_type}自适应分组结果：{[len(g) for g in groups]}个组")

        core_topic = texts[0].split("：")[0] if "：" in texts[0] else "未明确议题"

        for group_indices in groups:
            group_items = [items[valid_indices[idx]] for idx in group_indices]
            group_texts = [item["content"] for item in group_items]

            if len(group_texts) <= 1:
                integrated_item = {
                    "core_topic": group_items[0]["content"],
                    "related_items": group_items,
                    "has_conflict": False,
                    "conflict_details": None,
                    "conflict_probability": 0.0,
                    "reasoning": "组内观点数量不足"
                }
                integrated[section_type].append(integrated_item)
                continue

            conflict_analysis = enhanced_analyze_conflict_with_llm(group_texts, topic=core_topic)

            conflict_logs.append({
                "group_texts": group_texts,
                "analysis": conflict_analysis
            })

            group_embeddings = semantic_model.encode(group_texts, normalize_embeddings=True)
            group_similarity = np.mean(cosine_similarity(group_embeddings))
            dynamic_threshold = max(0.2, min(0.4, 0.4 - (0.5 - group_similarity) * 0.3))

            has_conflict = conflict_analysis["has_conflict"] and (
                conflict_analysis["probability"] >= dynamic_threshold or 
                (conflict_analysis["has_conflict"] and conflict_analysis["probability"] > 0.3)
            )

            integrated_item = {
                "core_topic": group_items[0]["content"],
                "related_items": group_items,
                "has_conflict": has_conflict,
                "conflict_details": conflict_analysis["details"],
                "conflict_probability": conflict_analysis["probability"],
                "reasoning": conflict_analysis["reasoning"],
                "dynamic_threshold": round(dynamic_threshold, 3)
            }
            integrated[section_type].append(integrated_item)

    os.makedirs(integrated_dir, exist_ok=True)
    with open(os.path.join(integrated_dir, "conflict_detection_logs.json"), "w", encoding="utf-8") as f:
        json.dump(conflict_logs, f, ensure_ascii=False, indent=2)

    return integrated


# -------------------------- 4. 整合文档生成模块 --------------------------
def generate_integrated_word(integrated_data, output_path):
    output_dir = os.path.dirname(output_path)
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    title = doc.add_heading("专家讨论整合报告（语义级冲突检测）", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"注：红色标注为冲突观点，冲突细节附后。")
    doc.add_paragraph("")

    for section_type in ["议题", "专家意见", "讨论结果"]:
        if section_type not in integrated_data or not integrated_data[section_type]:
            continue

        doc.add_heading(f"一、{section_type}", level=2)

        for idx, item in enumerate(integrated_data[section_type], 1):
            p = doc.add_paragraph()
            p.add_run(f"{idx}. 核心议题：").bold = True
            p.add_run(item["core_topic"])

            if item["has_conflict"]:
                conflict_p = doc.add_paragraph()
                conflict_p.add_run(f"[冲突提示]（概率：{item['conflict_probability']:.2f}，阈值：{item['dynamic_threshold']}）：").font.color.rgb = RGBColor(255, 0, 0)
                conflict_p.add_run(item["conflict_details"] or "未明确冲突点")
                conflict_p = doc.add_paragraph()
                conflict_p.add_run(f"判定依据：{item['reasoning']}").italic = True

            p = doc.add_paragraph("相关观点：")
            p.paragraph_format.left_indent = 20
            for related in item["related_items"]:
                rel_p = doc.add_paragraph()
                rel_p.paragraph_format.left_indent = 40
                rel_p.add_run(f"- {related['content']} ").font.color.rgb = RGBColor(255, 0, 0) if item["has_conflict"] else None
                rel_p.add_run(f"（来源：{related['file_id']}，{'; '.join(related['sources'])}）").italic = True
                if not related.get("valid_source", True):
                    rel_p.add_run(" [注意：来源格式不规范，可能影响冲突检测]").font.color.rgb = RGBColor(255, 165, 0)

        doc.add_paragraph("")

    doc.save(output_path)
    print(f"语义级整合Word文档已保存：{output_path}")


def generate_integrated_excel(integrated_data, output_path):
    output_dir = os.path.dirname(output_path)
    os.makedirs(output_dir, exist_ok=True)

    rows = []
    conflict_fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
    warning_fill = PatternFill(start_color="FFFFCC", end_color="FFFFCC", fill_type="solid")

    for section_type in ["议题", "专家意见", "讨论结果"]:
        if section_type not in integrated_data or not integrated_data[section_type]:
            continue

        for idx, item in enumerate(integrated_data[section_type], 1):
            rows.append({
                "类型": section_type,
                "序号": f"{idx}. 核心议题",
                "内容": item["core_topic"],
                "冲突标识": f"冲突（概率：{item['conflict_probability']:.2f}，阈值：{item['dynamic_threshold']}）：{item['conflict_details'] or '未明确冲突点'}" 
                            if item["has_conflict"] else "无冲突",
                "判定依据": item["reasoning"],
                "来源": "整合生成",
                "来源状态": "有效"
            })

            for rel_idx, related in enumerate(item["related_items"], 1):
                rows.append({
                    "类型": "",
                    "序号": f"{idx}.{rel_idx}",
                    "内容": related["content"],
                    "冲突标识": "冲突观点" if item["has_conflict"] else "无冲突",
                    "判定依据": "",
                    "来源": f"{related['file_id']}：{'; '.join(related['sources'])}",
                    "来源状态": "有效" if related.get("valid_source", True) else "格式不规范（可能影响冲突检测）"
                })

    df = pd.DataFrame(rows)
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='语义冲突检测')
        worksheet = writer.sheets['语义冲突检测']

        worksheet.column_dimensions['A'].width = 12
        worksheet.column_dimensions['B'].width = 10
        worksheet.column_dimensions['C'].width = 60
        worksheet.column_dimensions['D'].width = 50
        worksheet.column_dimensions['E'].width = 40
        worksheet.column_dimensions['F'].width = 40
        worksheet.column_dimensions['G'].width = 30

        bold_font = Font(bold=True)
        for cell in worksheet[1]:
            cell.font = bold_font
            cell.alignment = Alignment(horizontal='center')

        for row_idx in range(2, len(rows) + 2):
            if worksheet.cell(row=row_idx, column=4).value == "冲突观点":
                for col in range(1, 8):
                    worksheet.cell(row=row_idx, column=col).fill = conflict_fill
            if worksheet.cell(row=row_idx, column=7).value.startswith("格式不规范"):
                for col in range(1, 8):
                    worksheet.cell(row=row_idx, column=col).fill = warning_fill

        current_type = None
        merge_start = 2
        for row_idx in range(2, len(rows) + 2):
            cell_type = worksheet.cell(row=row_idx, column=1).value
            if cell_type != current_type or row_idx == len(rows) + 1:
                if current_type is not None:
                    worksheet.merge_cells(start_row=merge_start, start_column=1,
                                         end_row=row_idx - 1, end_column=1)
                current_type = cell_type
                merge_start = row_idx

    print(f"语义级整合Excel文档已保存：{output_path}")


# -------------------------- 模拟数据生成模块 --------------------------
class MockDataGenerator:
    @staticmethod
    def generate_mock_dialogue(num_turns=10, has_conflict=True):
        speakers = list(SPEAKER_MAPPING.values())
        topics = [
            "充电桩安装位置选择", 
            "电力负荷峰值控制", 
            "设备维护周期"
        ]
        opinions = {
            "充电桩安装位置选择": {
                "pro": [
                    "必须在停车场安装（电力容量充足且用户方便）",
                    "优先停车场，这里的供电线路无需改造"
                ],
                "anti": [
                    "绝对不能在停车场（会占用消防通道）",
                    "应强制在教学楼旁，覆盖师生步行人群"
                ]
            },
            "电力负荷峰值控制": {
                "pro": [
                    "必须限制高峰功率（8-20点）至50kW以下，否则会跳闸",
                    "高峰时段必须限流，电网改造费用太高"
                ],
                "anti": [
                    "完全无需限制功率，现有电网能承受（设计院评估过）",
                    "限制功率会导致充电速度太慢，用户会投诉"
                ]
            },
            "设备维护周期": {
                "pro": [
                    "必须每月维护一次，否则故障风险会增加30%",
                    "高频维护（每月）能延长设备寿命至5年以上"
                ],
                "anti": [
                    "每季度维护一次足够，每月维护成本增加2倍",
                    "设备稳定性高，过度维护是资源浪费"
                ]
            }
        }

        dialogue = []
        current_topic = random.choice(list(opinions.keys()))
        start_time = 0.0
        pro_used = False

        for i in range(num_turns):
            duration = random.uniform(3.0, 7.0)
            end_time = start_time + duration
            speaker = random.choice(speakers)

            if has_conflict and i > 2 and random.random() < 0.7 and pro_used:
                content = random.choice(opinions[current_topic]["anti"])
            else:
                if not pro_used:
                    content = random.choice(opinions[current_topic]["pro"])
                    pro_used = True
                else:
                    content = f"{random.choice(['补充一点，', '另外，', '而且，'])}" + random.choice(opinions[current_topic]["pro"])

            dialogue.append({
                "mapped_speaker": speaker,
                "start": round(start_time, 1),
                "end": round(end_time, 1),
                "text": f"{current_topic}：{content}"
            })
            start_time = end_time

        return dialogue, current_topic

    @staticmethod
    def generate_mock_ground_truth(dialogue):
        topics = [d["text"].split("：")[0] for d in dialogue]
        contents = [d["text"].split("：")[1] for d in dialogue]
        ground_truth = [False] * len(dialogue)

        conflict_pairs = {
            "充电桩安装位置选择": [("停车场", "教学楼"), ("必须", "绝对不能")],
            "电力负荷峰值控制": [("必须限制", "完全无需限制"), ("50kW以下", "太慢")],
            "设备维护周期": [("每月维护", "每季度维护"), ("必须", "足够")]
        }

        for i in range(len(contents)):
            for j in range(i+1, len(contents)):
                if topics[i] == topics[j] and topics[i] in conflict_pairs:
                    for (a, b) in conflict_pairs[topics[i]]:
                        if (a in contents[i] and b in contents[j]) or (b in contents[i] and a in contents[j]):
                            ground_truth[i] = True
                            ground_truth[j] = True
                            break

        return ground_truth


# -------------------------- 评估指标计算模块 --------------------------
class Evaluator:
    @staticmethod
    def evaluate_generation_quality(structured_data, ground_truth_dialogue, semantic_model=None):
        source_ids = set()
        for sec in structured_data.values():
            for item in sec:
                for src in item["sources"]:
                    ids = re.findall(r'\[(\d+)\]', src)
                    source_ids.update(ids)

        total_ids = set(str(i+1) for i in range(len(ground_truth_dialogue)))
        completeness = len(source_ids & total_ids) / len(total_ids) if total_ids else 0.0

        correct = 0
        total = 0

        ground_truth_texts = [d["text"].lower() for d in ground_truth_dialogue if d["text"].strip()]
        structured_texts = []
        for sec in structured_data.values():
            for item in sec:
                if item["content"].strip():
                    structured_texts.append(item["content"].lower())
                    total += 1

        if total == 0:
            accuracy = 0.0
        else:
            if semantic_model and len(ground_truth_texts) > 0 and len(structured_texts) > 0:
                gt_embeddings = semantic_model.encode(ground_truth_texts, normalize_embeddings=True)
                struct_embeddings = semantic_model.encode(structured_texts, normalize_embeddings=True)
                similarity_matrix = cosine_similarity(struct_embeddings, gt_embeddings)

                for sim_scores in similarity_matrix:
                    if max(sim_scores) >= GENERATION_ACCURACY_THRESHOLD:
                        correct += 1
            else:
                for struct_text in structured_texts:
                    struct_words = set(re.findall(r'\b\w+\b', struct_text)) - {'的', '是', '在', '应', '必须', '要', '有', '更'}
                    if not struct_words:
                        continue

                    for gt_text in ground_truth_texts:
                        gt_words = set(re.findall(r'\b\w+\b', gt_text))
                        if struct_words.issubset(gt_words):
                            correct += 1
                            break

        accuracy = correct / total if total else 0.0

        structured_ratio = 1.0 if all(k in structured_data for k in ["议题", "专家意见", "讨论结果"]) else 0.0
        if structured_ratio > 0:
            for sec in structured_data.values():
                for item in sec:
                    if "来源未知" in item["sources"] and not item.get("valid_source", False):
                        structured_ratio -= 0.05
            structured_ratio = max(structured_ratio, 0.0)

        return {
            "完整性": round(completeness, 3),
            "准确率": round(accuracy, 3),
            "结构化程度": round(structured_ratio, 3)
        }

    @staticmethod
    def evaluate_conflict_detection(integrated_data, ground_truth_labels):
        if sum(ground_truth_labels) == 0:
            return None

        pred_labels = []
        for sec in integrated_data.values():
            for item in sec:
                pred_labels.extend([item["has_conflict"]] * len(item["related_items"]))

        min_len = min(len(pred_labels), len(ground_truth_labels))
        pred_labels = pred_labels[:min_len]
        ground_truth_labels = ground_truth_labels[:min_len]

        precision = precision_score(ground_truth_labels, pred_labels, zero_division=0)
        recall = recall_score(ground_truth_labels, pred_labels, zero_division=0)
        f1 = f1_score(ground_truth_labels, pred_labels, zero_division=0)

        return {
            "准确率": round(precision, 3),
            "召回率": round(recall, 3),
            "F1值": round(f1, 3)
        }

    @staticmethod
    def evaluate_efficiency(func, *args, **kwargs):
        import timeit
        import psutil
        process = psutil.Process()
        start_time = timeit.default_timer()

        start_mem = process.memory_info().rss

        result = func(*args, **kwargs)

        end_mem = process.memory_info().rss

        time_cost = timeit.default_timer() - start_time
        mem_usage_mb = (end_mem - start_mem) / 1024 / 1024

        return {
            "处理时间(秒)": round(time_cost, 3),
            "内存占用(MB)": round(mem_usage_mb, 3)
        }, result


# -------------------------- 系统流程封装 --------------------------
class ConflictDetectionSystem:
    def __init__(self):
        self.asr_model = None
        self.speaker_diarization = None
        self.semantic_model = None
        self._init_models()

    def _init_models(self):
        print("===== 初始化系统模型 =====")
        # 1. Whisper模型（引用config中的模型大小）
        self.asr_model = whisper.load_model(WHISPER_MODEL_SIZE, device=device)
        print(f"Whisper模型加载完成（设备：{device}）")

        # 2. 说话人分离模型 - 直接从缓存加载（离线模式）
        try:
            print("从本地缓存加载说话人分离模型: pyannote/speaker-diarization-3.1")
            self.speaker_diarization = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1"
            ).to(torch.device(device))
            print("说话人分离模型加载完成")
        except Exception as e:
            print(f"说话人分离模型加载失败：{e}")
            print("提示：请确保模型已正确复制到缓存目录 /root/autodl-tmp/huggingface/hub/，并设置了 HF_HUB_OFFLINE=1")
            raise

        # 3. 语义模型（从本地路径）
        try:
            if not os.path.exists(SEMANTIC_MODEL_PATH):
                raise FileNotFoundError(f"语义模型路径不存在：{SEMANTIC_MODEL_PATH}")

            self.semantic_model = SentenceTransformer(
                SEMANTIC_MODEL_PATH,
                local_files_only=True
            )
            self.semantic_model.to(device)
            print(f"语义模型加载完成（路径：{SEMANTIC_MODEL_PATH}，设备：{device}）")
        except Exception as e:
            print(f"语义模型加载失败：{e}")
            print("提示：如果模型下载失败，可以手动从 https://hf-mirror.com/sentence-transformers/all-MiniLM-L6-v2 下载")
            raise

    def process_audio(self, audio_path=None, mock_dialogue=None):
        if mock_dialogue is not None:
            raw_dialogue = mock_dialogue
            print("使用模拟对话数据进行处理")
        else:
            print(f"开始转写音频：{audio_path}")
            asr_result = self.asr_model.transcribe(audio_path, fp16=False)
            print("转写完成，开始分离说话人...")
            diarization_result = self.speaker_diarization(audio_path)
            final_result = diarize_text(asr_result, diarization_result)
            raw_dialogue = [
                {
                    'mapped_speaker': mapped_spk,
                    'start': round(start, 1),
                    'end': round(end, 1),
                    'text': sent
                } 
                for seg, raw_spk, mapped_spk, sent, start, end in final_result
            ]

        print("生成结构化分析...")
        llm_input = format_dialogue_for_llm(raw_dialogue)
        structured_text = generate_structured_with_llm(llm_input)
        if "生成失败" in structured_text:
            print(f"警告：{structured_text}")

        sections = parse_llm_output(structured_text, file_id="test", raw_dialogue=raw_dialogue)
        print(f"结构化完成：议题{len(sections['议题'])}条，意见{len(sections['专家意见'])}条")

        print("开始冲突检测...")
        integrated_data = optimized_detect_semantic_conflicts(sections, self.semantic_model)

        return raw_dialogue, sections, integrated_data

    def run_evaluation(self, num_samples=5):
        print("\n===== 开始模拟数据评估 =====")
        all_metrics = {
            "生成质量": [],
            "冲突检测性能": [],
            "处理效率": []
        }

        valid_conflict_samples = 0

        for i in range(num_samples):
            print(f"\n----- 评估样本 {i+1}/{num_samples} -----")
            mock_dialogue, _ = MockDataGenerator.generate_mock_dialogue(
                num_turns=random.randint(8, 15), 
                has_conflict=True
            )
            ground_truth_labels = MockDataGenerator.generate_mock_ground_truth(mock_dialogue)
            conflict_count = sum(ground_truth_labels)

            print(f"样本 {i+1} 真实冲突数量：{conflict_count}")
            if conflict_count == 0:
                print("警告：该样本无真实冲突，跳过冲突检测指标计算")
            else:
                valid_conflict_samples += 1

            efficiency_metrics, (_, structured_data, integrated_data) = Evaluator.evaluate_efficiency(
                self.process_audio, 
                mock_dialogue=mock_dialogue
            )
            all_metrics["处理效率"].append(efficiency_metrics)

            gen_metrics = Evaluator.evaluate_generation_quality(
                structured_data, 
                mock_dialogue,
                semantic_model=self.semantic_model
            )
            all_metrics["生成质量"].append(gen_metrics)
            print(f"生成质量：{gen_metrics}")

            conflict_metrics = Evaluator.evaluate_conflict_detection(integrated_data, ground_truth_labels)
            if conflict_metrics is not None:
                all_metrics["冲突检测性能"].append(conflict_metrics)
                print(f"冲突检测性能：{conflict_metrics}")
            else:
                print("冲突检测性能：跳过（无真实冲突）")

            print(f"处理效率：{efficiency_metrics}")

        avg_metrics = {
            "生成质量": {
                "平均完整性": round(sum(m["完整性"] for m in all_metrics["生成质量"])/num_samples, 3),
                "平均准确率": round(sum(m["准确率"] for m in all_metrics["生成质量"])/num_samples, 3),
                "平均结构化程度": round(sum(m["结构化程度"] for m in all_metrics["生成质量"])/num_samples, 3)
            },
            "处理效率": {
                "平均处理时间(秒)": round(sum(m["处理时间(秒)"] for m in all_metrics["处理效率"])/num_samples, 3),
                "平均内存占用(MB)": round(sum(m["内存占用(MB)"] for m in all_metrics["处理效率"])/num_samples, 3)
            }
        }

        if valid_conflict_samples > 0:
            avg_metrics["冲突检测性能"] = {
                "平均准确率": round(sum(m["准确率"] for m in all_metrics["冲突检测性能"])/valid_conflict_samples, 3),
                "平均召回率": round(sum(m["召回率"] for m in all_metrics["冲突检测性能"])/valid_conflict_samples, 3),
                "平均F1值": round(sum(m["F1值"] for m in all_metrics["冲突检测性能"])/valid_conflict_samples, 3)
            }
        else:
            avg_metrics["冲突检测性能"] = "无有效冲突样本（所有样本均无冲突）"

        self._visualize_metrics(avg_metrics)
        return avg_metrics

    def _visualize_metrics(self, avg_metrics):
        import matplotlib.pyplot as plt
        plt.rcParams["font.family"] = ["Arial", "Helvetica", "sans-serif"]
        plt.rcParams["axes.unicode_minus"] = False

        has_conflict_metrics = isinstance(avg_metrics.get("冲突检测性能"), dict)
        fig, axes = plt.subplots(1, 3 if has_conflict_metrics else 2, figsize=(18 if has_conflict_metrics else 12, 5))
        axes = [axes] if not isinstance(axes, np.ndarray) else axes

        gen_keys = ["Completeness", "Accuracy", "Structured Degree"]
        gen_values = [
            avg_metrics["生成质量"]["平均完整性"],
            avg_metrics["生成质量"]["平均准确率"],
            avg_metrics["生成质量"]["平均结构化程度"]
        ]
        axes[0].bar(gen_keys, gen_values, color='skyblue')
        axes[0].set_title("Generation Quality Metrics")
        axes[0].set_ylim(0, 1.0)
        for i, v in enumerate(gen_values):
            axes[0].text(i, v+0.02, f"{v:.3f}", ha='center')

        if has_conflict_metrics:
            conf_keys = ["Precision", "Recall", "F1 Score"]
            conf_values = [
                avg_metrics["冲突检测性能"]["平均准确率"],
                avg_metrics["冲突检测性能"]["平均召回率"],
                avg_metrics["冲突检测性能"]["平均F1值"]
            ]
            axes[1].bar(conf_keys, conf_values, color='lightgreen')
            axes[1].set_title("Conflict Detection Performance")
            axes[1].set_ylim(0, 1.0)
            for i, v in enumerate(conf_values):
                axes[1].text(i, v+0.02, f"{v:.3f}", ha='center')

            eff_ax_idx = 2
        else:
            eff_ax_idx = 1

        eff_keys = ["Processing Time (s)", "Memory Usage (MB)"]
        eff_values = [
            avg_metrics["处理效率"]["平均处理时间(秒)"],
            avg_metrics["处理效率"]["平均内存占用(MB)"]
        ]
        axes[eff_ax_idx].bar(eff_keys, eff_values, color='salmon')
        axes[eff_ax_idx].set_title("Processing Efficiency")
        for i, v in enumerate(eff_values):
            axes[eff_ax_idx].text(i, v+0.1, f"{v:.3f}", ha='center')

        plt.tight_layout()
        plt.savefig(os.path.join(integrated_dir, "evaluation_metrics.png"))
        print(f"\n评估指标可视化已保存：{os.path.join(integrated_dir, 'evaluation_metrics.png')}")
        plt.show()


# -------------------------- 主程序 --------------------------
def main():
    import timeit
    import psutil
    os.makedirs(audio_dir, exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(integrated_dir, exist_ok=True)
    
    print("当前工作目录:", os.getcwd())

    system = ConflictDetectionSystem()

    mode = input("请选择运行模式（1=处理真实音频，2=模拟数据评估）：")

    if mode == "1":
        wav_files = [f for f in os.listdir(audio_dir) if f.lower().endswith(".wav")]
        if not wav_files:
            print(f"未找到{audio_dir}中的WAV文件，请检查目录是否正确")
            return

        for file in wav_files:
            audio_path = os.path.join(audio_dir, file)
            print(f"\n----- 处理真实音频：{file} -----")
            raw_dialogue, _, integrated_data = system.process_audio(audio_path=audio_path)

            base_name = os.path.splitext(file)[0]
            integrated_word_path = os.path.join(integrated_dir, f"{base_name}_report.docx")
            integrated_excel_path = os.path.join(integrated_dir, f"{base_name}_report.xlsx")
            generate_integrated_word(integrated_data, integrated_word_path)
            generate_integrated_excel(integrated_data, integrated_excel_path)

        print("\n===== 真实音频处理完成 =====")

    elif mode == "2":
        try:
            num_samples = int(input("请输入评估样本数量（建议3-5）："))
            if num_samples < 1:
                raise ValueError
        except ValueError:
            print("无效输入，使用默认样本数量3")
            num_samples = DEFAULT_EVAL_SAMPLES

        avg_metrics = system.run_evaluation(num_samples=num_samples)
        print("\n===== 模拟评估完成 =====")
        print("平均评估指标：")
        for metric_type, metrics in avg_metrics.items():
            print(f"\n{metric_type}：")
            if isinstance(metrics, dict):
                for k, v in metrics.items():
                    print(f"  {k}：{v}")
            else:
                print(f"  {metrics}")


if __name__ == "__main__":
    main()
